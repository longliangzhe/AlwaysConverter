#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文档转换器
支持PDF、DOC、DOCX、TXT、RTF、ODT等文档格式转换
"""

import os
from core.base_converter import BaseConverter
from utils.logger import get_logger

logger = get_logger(__name__)

try:
    from docx import Document
    from docx.shared import Inches
    import PyPDF2
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    HAS_DOCUMENT_LIBS = True
except ImportError:
    HAS_DOCUMENT_LIBS = False
    logger.warning("文档处理库未安装，文档转换功能将受限")


class DocumentConverter(BaseConverter):
    """文档转换器"""
    
    def __init__(self, config):
        super().__init__(config)
        self.supported = HAS_DOCUMENT_LIBS
    
    def convert(self, input_path: str, output_path: str, source_format: str, target_format: str) -> bool:
        """
        执行文档转换
        
        Args:
            input_path: 输入文件路径
            output_path: 输出文件路径
            source_format: 源文件格式
            target_format: 目标文件格式
            
        Returns:
            转换是否成功
        """
        if not self.supported:
            logger.error("缺少必要的文档处理库")
            return False
        
        try:
            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            # 根据源格式和目标格式选择转换方法
            if source_format == 'pdf' and target_format == 'txt':
                return self._pdf_to_txt(input_path, output_path)
            elif source_format in ['docx', 'doc'] and target_format == 'txt':
                return self._docx_to_txt(input_path, output_path)
            elif source_format == 'txt' and target_format == 'docx':
                return self._txt_to_docx(input_path, output_path)
            elif source_format == 'txt' and target_format == 'pdf':
                return self._txt_to_pdf(input_path, output_path)
            else:
                logger.warning(f"不支持的文档转换: {source_format} -> {target_format}")
                return False
                
        except Exception as e:
            logger.error(f"文档转换失败: {str(e)}")
            return False
    
    def _pdf_to_txt(self, input_path: str, output_path: str) -> bool:
        """PDF转TXT"""
        try:
            with open(input_path, 'rb') as pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                with open(output_path, 'w', encoding='utf-8') as txt_file:
                    txt_file.write(text)
            
            logger.info(f"PDF转TXT成功: {input_path} -> {output_path}")
            return True
        except Exception as e:
            logger.error(f"PDF转TXT失败: {str(e)}")
            return False
    
    def _docx_to_txt(self, input_path: str, output_path: str) -> bool:
        """DOCX转TXT"""
        try:
            doc = Document(input_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            with open(output_path, 'w', encoding='utf-8') as txt_file:
                txt_file.write(text)
            
            logger.info(f"DOCX转TXT成功: {input_path} -> {output_path}")
            return True
        except Exception as e:
            logger.error(f"DOCX转TXT失败: {str(e)}")
            return False
    
    def _txt_to_docx(self, input_path: str, output_path: str) -> bool:
        """TXT转DOCX"""
        try:
            doc = Document()
            
            with open(input_path, 'r', encoding='utf-8') as txt_file:
                for line in txt_file:
                    doc.add_paragraph(line.rstrip())
            
            doc.save(output_path)
            
            logger.info(f"TXT转DOCX成功: {input_path} -> {output_path}")
            return True
        except Exception as e:
            logger.error(f"TXT转DOCX失败: {str(e)}")
            return False
    
    def _txt_to_pdf(self, input_path: str, output_path: str) -> bool:
        """TXT转PDF"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            
            with open(input_path, 'r', encoding='utf-8') as txt_file:
                text = txt_file.read()
            
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # 简单的文本绘制
            text_object = c.beginText(50, height - 50)
            text_object.setFont("Helvetica", 12)
            
            for line in text.split('\n'):
                text_object.textLine(line)
            
            c.drawText(text_object)
            c.save()
            
            logger.info(f"TXT转PDF成功: {input_path} -> {output_path}")
            return True
        except Exception as e:
            logger.error(f"TXT转PDF失败: {str(e)}")
            return False